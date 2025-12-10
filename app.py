from flask import Flask, render_template, request, session, send_file
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import io
import os

app = Flask(__name__)
app.secret_key = 'your-secret-key-here-change-this'  # Change this to a random string

# Configure Gemini API
GEMINI_API_KEY = "AIzaSyAUznpVqGIjGu7qinINqNHBg10WXnXbhUw"
genai.configure(api_key=GEMINI_API_KEY)

# Font mapping for documents
FONT_MAP = {
    'inter': 'Calibri',
    'roboto': 'Arial',
    'opensans': 'Calibri',
    'lora': 'Georgia',
    'montserrat': 'Arial'
}

@app.route("/", methods=["GET", "POST"])
def index():
    tailored_resume = ""
    cover_letter = ""
    bullet_variants = ""
    error_message = ""
    selected_font = "inter"
    
    if request.method == "POST":
        resume = request.form["resume"]
        job_description = request.form["job_description"]
        selected_font = request.form.get("selected_font", "inter")
        
        try:
            # Generate AI-powered outputs
            tailored_resume, cover_letter, bullet_variants = generate_tailored_content(
                resume, job_description
            )
            
            # Store in session for downloads
            session['tailored_resume'] = tailored_resume
            session['cover_letter'] = cover_letter
            session['bullet_variants'] = bullet_variants
            session['selected_font'] = selected_font
            
        except Exception as e:
            error_message = f"❌ Error: {str(e)}"
            print(f"Full error details: {e}")
    
    return render_template(
        "index.html",
        tailored_resume=tailored_resume,
        cover_letter=cover_letter,
        bullet_variants=bullet_variants,
        error_message=error_message,
        selected_font=selected_font
    )

@app.route("/download/<content_type>/<format_type>")
def download(content_type, format_type):
    """
    Download route for PDF and Word documents
    content_type: 'resume', 'cover', 'bullets'
    format_type: 'word', 'pdf'
    """
    # Get content from session
    content_map = {
        'resume': session.get('tailored_resume', ''),
        'cover': session.get('cover_letter', ''),
        'bullets': session.get('bullet_variants', '')
    }
    
    content = content_map.get(content_type, '')
    if not content:
        return "No content available", 404
    
    selected_font = session.get('selected_font', 'inter')
    
    # Generate filename
    filename_map = {
        'resume': 'Tailored_Resume',
        'cover': 'Cover_Letter',
        'bullets': 'Bullet_Point_Variants'
    }
    filename = filename_map.get(content_type, 'Document')
    
    if format_type == 'word':
        return generate_word(content, filename, selected_font)
    elif format_type == 'pdf':
        return generate_pdf(content, filename, selected_font)
    
    return "Invalid format", 400

def generate_word(content, filename, font_name):
    """Generate a Word document"""
    doc = Document()
    
    # Set font
    font = FONT_MAP.get(font_name, 'Calibri')
    
    # Add content
    for line in content.split('\n'):
        if line.strip():
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.name = font
                run.font.size = Pt(11)
        else:
            doc.add_paragraph()  # Empty line
    
    # Save to BytesIO
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return send_file(
        file_stream,
        as_attachment=True,
        download_name=f'{filename}.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

def generate_pdf(content, filename, font_name):
    """Generate a PDF document"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter,
                           rightMargin=72, leftMargin=72,
                           topMargin=72, bottomMargin=18)
    
    # Container for the 'Flowable' objects
    elements = []
    
    # Define styles
    styles = getSampleStyleSheet()
    style = ParagraphStyle(
        'CustomStyle',
        parent=styles['Normal'],
        fontSize=11,
        leading=14,
        spaceAfter=6,
    )
    
    # Add content
    for line in content.split('\n'):
        if line.strip():
            # Replace special characters
            line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            p = Paragraph(line, style)
            elements.append(p)
        else:
            elements.append(Spacer(1, 0.2*inch))
    
    # Build PDF
    doc.build(elements)
    buffer.seek(0)
    
    return send_file(
        buffer,
        as_attachment=True,
        download_name=f'{filename}.pdf',
        mimetype='application/pdf'
    )

def generate_tailored_content(resume, job_description):
    """
    Uses Gemini AI to automatically tailor the resume to match the job description.
    Returns: (tailored_resume, cover_letter, bullet_variants)
    """
    
    # Create the AI prompt
    prompt = f"""You are an expert resume writer and career coach. Your job is to help job seekers get interviews.

I will provide you with:
1. A MASTER RESUME (the candidate's full experience)
2. A JOB DESCRIPTION (the role they're applying for)

Your task is to create THREE outputs:

**OUTPUT 1 - TAILORED RESUME:**
- Rewrite the resume to perfectly match the job description
- Use keywords and phrases from the job description naturally
- Keep all real experience but rephrase it to align with the job requirements
- Remove or minimize irrelevant experience that doesn't match
- Add skills mentioned in the job description if the candidate likely has them
- Make it ATS-friendly (Applicant Tracking System optimized)
- Keep it professional and honest (no lies, just better positioning)

**OUTPUT 2 - COVER LETTER:**
- Write a compelling, professional cover letter
- Reference specific requirements from the job description
- Highlight how the candidate's experience matches
- Show enthusiasm for the role and company
- Keep it to 3-4 paragraphs maximum

**OUTPUT 3 - BULLET POINT VARIANTS:**
- For each major role/experience in the resume, provide 3 alternative bullet point options
- Use the format: Achievement + Impact + Metric (when possible)
- Each variant should emphasize different aspects of the same experience

---

**MASTER RESUME:**
{resume}

---

**JOB DESCRIPTION:**
{job_description}

---

**FORMAT YOUR RESPONSE EXACTLY LIKE THIS:**

===TAILORED_RESUME===
[Your tailored resume here]

===COVER_LETTER===
[Your cover letter here]

===BULLET_VARIANTS===
[Your bullet point variants here]
"""

    # Use Gemini 2.5 Flash - stable and fast
    model = genai.GenerativeModel('models/gemini-2.5-flash')
    response = model.generate_content(prompt)
    
    # Get the response text
    full_response = response.text
    
    # Split into sections
    tailored_resume = extract_section(full_response, "===TAILORED_RESUME===", "===COVER_LETTER===")
    cover_letter = extract_section(full_response, "===COVER_LETTER===", "===BULLET_VARIANTS===")
    bullet_variants = extract_section(full_response, "===BULLET_VARIANTS===", "===END===")
    
    return tailored_resume, cover_letter, bullet_variants

def extract_section(text, start_marker, end_marker):
    """
    Extracts content between two markers.
    """
    try:
        start_idx = text.find(start_marker)
        if start_idx == -1:
            return "Section not found"
        
        start_idx += len(start_marker)
        
        end_idx = text.find(end_marker, start_idx)
        if end_idx == -1:
            # If no end marker, take everything after start
            return text[start_idx:].strip()
        
        return text[start_idx:end_idx].strip()
    except Exception as e:
        return f"Error extracting section: {str(e)}"

if __name__ == "__main__":
    print("🚀 Starting AI Resume Tailoring Assistant...")
    print(f"✅ Gemini API configured")
    print("📍 Open http://127.0.0.1:5000 in your browser")
    if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=False)